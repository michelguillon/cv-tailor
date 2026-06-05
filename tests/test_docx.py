"""--docx stretch (SPEC §"Stretch"): harvest a source CV's formatting conventions and
render the assembled clean CV as a styled .docx. Fully deterministic — a fixture .docx
is built in-process, so there's no real CV and no API call (tests/CLAUDE.md)."""

from docx import Document
from docx.shared import Pt

from tailor.phases import phase6_docx
from tailor.phases.phase6_docx import DocxConventions


def _build_source_cv(path):
    """A tiny stand-in for a real CV: a 22pt bold name, 13pt bold heading, 10pt body —
    all in Garamond — so the harvester has clear modal/size signals to read back."""
    doc = Document()

    def para(text, *, size, bold=False, font="Garamond"):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.name = font
        r.bold = bold
        return p

    para("Jane Doe", size=22, bold=True)              # name = largest text
    para("jane@example.com", size=10)
    para("Experience", size=13, bold=True)            # heading = the band below the name
    para("Senior Engineer", size=10, bold=True)       # role line
    para("Built systems", size=10)                    # body (modal size)
    para("Led delivery", size=10)
    para("Shipped product", size=10)
    doc.save(str(path))
    return path


def test_harvest_conventions_reads_sizes_and_font(tmp_path):
    src = _build_source_cv(tmp_path / "CV_source.docx")
    c = phase6_docx.harvest_conventions(src)
    assert c.body_size == 10.0          # modal body size
    assert c.name_size == 22.0          # largest text = the name
    assert c.heading_size == 13.0       # the distinct size between body and name
    assert c.heading_bold is True
    assert c.body_font == "Garamond"


CV_MD = (
    "Jane Doe\n"
    "jane@example.com | London\n"
    "\n"
    "## Experience — Acme\n"
    "\n"
    "**Senior Engineer | 2020–2024**\n"
    "\n"
    "- Built **scalable** systems\n"
    "- Led the team\n"
    "\n"
    "## Skills\n"
    "\n"
    "Python, Go, Rust\n"
)


def _by_text(doc):
    return {p.text: p for p in doc.paragraphs if p.text}


def test_render_applies_conventions_and_strips_markdown(tmp_path):
    conv = DocxConventions(body_font="Garamond", body_size=10, name_size=22,
                           heading_size=13, heading_bold=True)
    out = phase6_docx.render_cv_docx(CV_MD, conv, tmp_path / "cv_final.docx")
    doc = Document(str(out))
    paras = _by_text(doc)

    # Normal style carries the harvested body font/size.
    assert doc.styles["Normal"].font.name == "Garamond"
    assert doc.styles["Normal"].font.size == Pt(10)

    # Name: first block, bold, name size.
    name = doc.paragraphs[0]
    assert name.text == "Jane Doe"
    assert all(r.bold for r in name.runs) and name.runs[0].font.size == Pt(22)

    # Heading: bold at heading size, "## " stripped.
    head = paras["Experience — Acme"]
    assert head.runs[0].bold and head.runs[0].font.size == Pt(13)

    # Role line: the standalone **…** becomes a bold run with the markers stripped.
    role = paras["Senior Engineer | 2020–2024"]
    assert role.runs and all(r.bold for r in role.runs)

    # Bullet: real Word list style + inline **bold** rendered as a bold run.
    bullet = paras["Built scalable systems"]
    assert bullet.style.name == "List Bullet"
    assert any(r.bold and r.text == "scalable" for r in bullet.runs)

    # Plain body paragraph survives.
    assert "Python, Go, Rust" in paras


def test_write_cv_docx_end_to_end_harvest_then_render(tmp_path):
    src = _build_source_cv(tmp_path / "CV_source.docx")
    out = phase6_docx.write_cv_docx(CV_MD, src, tmp_path / "cv_final.docx")
    doc = Document(str(out))
    # Harvested Garamond/10 flow through to the rendered document.
    assert doc.styles["Normal"].font.name == "Garamond"
    assert doc.paragraphs[0].text == "Jane Doe"
    assert doc.paragraphs[0].runs[0].font.size == Pt(22)


def test_resolve_template_prefers_dominant_source(tmp_path):
    cvs = tmp_path / "cvs"
    cvs.mkdir()
    (cvs / "CV_Michel_Guillon_2026_AI.docx").write_bytes(b"x")
    (cvs / "CV_Michel_Guillon_2026_generic.docx").write_bytes(b"x")
    manifest = {
        "header": {"static": True, "source_cv": "AI"},
        "profile": {"static": False, "source_cv": "generic"},
        "skills": {"static": False, "source_cv": "generic"},
    }
    chosen = phase6_docx.resolve_template(manifest, cvs_dir=cvs)
    assert chosen.name == "CV_Michel_Guillon_2026_generic.docx"   # backs the most non-static sections


def test_resolve_template_none_when_no_cvs(tmp_path):
    assert phase6_docx.resolve_template({}, cvs_dir=tmp_path) is None
