"""Phase 6 (stretch) — clean CV as a .docx, in the source CV's formatting (SPEC §"Stretch").

`--docx` only. Renders the SAME assembled markdown as `cv_final.md` (so the two are
always in parity) into a Word document, applying formatting *conventions* harvested
from a source CV in the corpus — not a pixel clone (the tailored CV reorders and mixes
sections from several CVs, so an in-place edit of one source can't represent it).

Two steps:
  1. `harvest_conventions(source_docx)` — read the source .docx (reusing the table-aware
     `corpus.docx_loader`, since these CVs keep their content in a single table) and infer
     body font/size, the name size (largest text), the heading size (the band between
     body and name), and whether headings are bold.
  2. `render_cv_docx(cv_md, conventions, out)` — parse the assembled markdown
     (`## heading`, standalone `**role line**` (F-29), `- bullet`, plain paragraph; the
     leading block before the first heading is the name + contact) and emit styled
     paragraphs. Inline `**bold**` becomes bold runs.

Deterministic and provider-free; unit-tested against a fixture .docx (no real CV needed).
"""

from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt

from corpus.docx_loader import load_docx

__all__ = ["DocxConventions", "harvest_conventions", "render_cv_docx",
           "write_cv_docx", "resolve_template"]

DEFAULT_FONT = "Calibri"
DEFAULT_BODY_PT = 11.0
_BOLD = re.compile(r"\*\*(.+?)\*\*")


@dataclass
class DocxConventions:
    """The formatting signals harvested from a source CV (all sizes in points)."""
    body_font: str = DEFAULT_FONT
    body_size: float = DEFAULT_BODY_PT
    name_size: float = DEFAULT_BODY_PT * 1.7
    heading_size: float = DEFAULT_BODY_PT + 1.5
    heading_bold: bool = True
    role_bold: bool = True


def _modal(values):
    vals = [v for v in values if v is not None]
    return Counter(vals).most_common(1)[0][0] if vals else None


# --------------------------------------------------------------------------- #
# Harvest                                                                      #
# --------------------------------------------------------------------------- #

def _harvest_font(source_docx) -> str:
    """The modal run font across body + table cells (CV content lives in a table).
    Falls back to the Normal style font, then a sane default."""
    doc = Document(str(source_docx))
    names: list[str] = []

    def scan(paragraphs):
        for p in paragraphs:
            names.extend(r.font.name for r in p.runs if r.font.name)

    scan(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan(cell.paragraphs)
    if (modal := _modal(names)):
        return modal
    try:
        if (nm := doc.styles["Normal"].font.name):
            return nm
    except KeyError:
        pass
    return DEFAULT_FONT


def harvest_conventions(source_docx) -> DocxConventions:
    """Infer body/name/heading sizes + body font from a source CV (table-aware)."""
    paras = load_docx(source_docx)
    sizes = [p.rendered_size for p in paras if p.rendered_size]
    # Body = the modal size among bulleted lines (a CV body is mostly bullets); if the
    # source has no bullets, the modal size overall (body text dominates a CV).
    bullet_sizes = [p.rendered_size for p in paras if p.has_num_pr and p.rendered_size]
    body_size = _modal(bullet_sizes) or _modal(sizes) or DEFAULT_BODY_PT
    name_size = max(sizes) if sizes else body_size * 1.7
    # Headings: the largest distinct size strictly between body and the name.
    head_sizes = sorted({s for s in sizes if body_size < s < name_size}, reverse=True)
    heading_size = head_sizes[0] if head_sizes else round(body_size + 1.5, 1)
    heading_bold = (any(p.is_bold for p in paras if p.rendered_size == heading_size)
                    if head_sizes else True)
    return DocxConventions(
        body_font=_harvest_font(source_docx),
        body_size=float(body_size), name_size=float(name_size),
        heading_size=float(heading_size), heading_bold=bool(heading_bold),
    )


# --------------------------------------------------------------------------- #
# Render                                                                       #
# --------------------------------------------------------------------------- #

def _add_runs(paragraph, text: str, *, bold=False, size=None, font=None) -> None:
    """Add `text` to a paragraph, turning markdown **bold** into bold runs. `bold`
    forces the whole paragraph bold (name / heading / role line)."""
    if not text:
        return
    for i, part in enumerate(_BOLD.split(text)):     # odd indices are the **captured** spans
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = bool(bold or i % 2 == 1)
        if size is not None:
            run.font.size = Pt(size)
        if font:
            run.font.name = font


def render_cv_docx(cv_md: str, conventions: DocxConventions, out_path) -> Path:
    """Render assembled CV markdown to a styled .docx using harvested conventions."""
    c = conventions
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = c.body_font
    normal.size = Pt(c.body_size)

    seen_heading = False
    name_done = False
    for raw in cv_md.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if line.startswith("## "):                   # section / company heading
            seen_heading = True
            _add_runs(doc.add_paragraph(), line[3:].strip(),
                      bold=c.heading_bold, size=c.heading_size, font=c.body_font)
        elif not stripped:
            continue
        elif not seen_heading:                        # leading block: name, then contact
            if not name_done:
                _add_runs(doc.add_paragraph(), stripped, bold=True, size=c.name_size, font=c.body_font)
                name_done = True
            else:
                _add_runs(doc.add_paragraph(), stripped, size=c.body_size, font=c.body_font)
        elif stripped.startswith(("- ", "* ")):       # bullet
            _add_runs(doc.add_paragraph(style="List Bullet"), stripped[2:].strip(),
                      size=c.body_size, font=c.body_font)
        elif stripped.startswith("**") and stripped.endswith("**"):   # role/date line (F-29)
            _add_runs(doc.add_paragraph(), stripped, bold=c.role_bold, size=c.body_size, font=c.body_font)
        else:                                          # plain paragraph
            _add_runs(doc.add_paragraph(), stripped, size=c.body_size, font=c.body_font)

    out_path = Path(out_path)
    doc.save(str(out_path))
    return out_path


def write_cv_docx(cv_md: str, source_docx, out_path) -> Path:
    """Harvest conventions from `source_docx`, render `cv_md` to `out_path`."""
    return render_cv_docx(cv_md, harvest_conventions(source_docx), out_path)


def resolve_template(manifest: dict, cvs_dir: str | Path = "data/cvs") -> Path | None:
    """Pick the source .docx whose formatting to mirror: the corpus file backing the
    most non-static sections (all CVs are the same family, so any is acceptable — this
    just prefers the dominant source). None if the corpus dir has no .docx."""
    files = sorted(Path(cvs_dir).glob("*.docx"))
    if not files:
        return None
    shorts = Counter(m.get("source_cv") for m in manifest.values()
                     if not m.get("static") and m.get("source_cv"))
    for short, _ in shorts.most_common():
        for f in files:
            if short and short in f.stem:
                return f
    return files[0]
