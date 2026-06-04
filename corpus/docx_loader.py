"""corpus/docx_loader.py — flatten a .docx into a list of Paragraph objects.

Adapted from the Week 1 RAG pipeline's `loaders/docx_loader.py` (reuse, not
reinvention). The parsing logic is unchanged — same style-inheritance size walk,
same numPr bullet detection, same table date-column pairing — because the CV
corpus is the same family of documents that logic was built for.

Two deliberate simplifications for cv-tailor (recorded in LEARNING_NOTES F-04):
  - cv-tailor ingests only .docx (no PDF), so the format-agnostic `models/`
    split is collapsed: the `Paragraph` dataclass lives here, in the one loader.
  - `source_format` is dropped — nothing branches on it when there is one format.

KEY INSIGHT (carried over): read every cell of a row, not just the first.
These CVs use a single table; most rows are one merged cell, but role and
education rows carry a SECOND column holding dates. Reading only cells[0]
silently drops every date. We iterate the raw <w:tc> children: first cell is
content, last cell is the date column, paired with content paragraphs by index.

KEY INSIGHT (carried over): detect bullets by the numPr element, not the style
name. In this corpus bullets appear under both 'List Paragraph' and 'Normal'.

KEY INSIGHT (carried over): report the *rendered* font size — resolve the
run-override → paragraph-style → base-style chain — because section headers and
company names here are sized, not heading-styled.
"""

from dataclasses import dataclass

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell

__all__ = ["Paragraph", "load_docx"]


@dataclass
class Paragraph:
    """One paragraph of a .docx, with the signals needed to discover structure.

    Fields:
      text          — plain text (stripped)
      style_name    — paragraph style ("Heading 1", "Normal", ...) or "(none)"
      rendered_size — font size in points after the inheritance chain, or None
      is_bold       — True if the paragraph renders bold
      has_num_pr    — True if it is a list item (a <w:numPr> element)
      in_table      — True if it came from inside a table
      date          — date string paired from a sibling table cell, else ""
      override      — True if a direct run-level font-size override is present
                      (diagnostic: heading level is not a reliable size proxy)
    """

    text: str
    style_name: str | None
    rendered_size: float | None
    is_bold: bool
    has_num_pr: bool
    in_table: bool
    date: str = ""
    override: bool = False

    def __repr__(self) -> str:
        size = "?" if self.rendered_size is None else f"{self.rendered_size:g}pt"
        style = self.style_name or "-"
        preview = self.text[:60] + ("…" if len(self.text) > 60 else "")
        bold = " B" if self.is_bold else "  "
        bullet = "•" if self.has_num_pr else " "
        return f"Paragraph({style!r:20} {size:>6}{bold} {bullet} {preview!r})"


def effective_size(para):
    """Rendered font size in points, or None.

    Report the size the reader actually sees: a direct run override wins,
    otherwise walk the paragraph-style → base-style chain.
    """
    sizes = [r.font.size.pt for r in para.runs if r.font.size is not None]
    if sizes:
        return max(sizes)
    style = para.style
    seen = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        if style.font is not None and style.font.size is not None:
            return style.font.size.pt
        style = style.base_style
    return None


def has_numbering(para):
    """True if the paragraph carries a numPr element (bulleted/numbered list)."""
    pPr = para._p.pPr
    return pPr is not None and pPr.find(qn("w:numPr")) is not None


def _date_cell_texts(tc):
    """Plain text of each non-empty <w:p> in a raw <w:tc> (the date column)."""
    texts = []
    for p in tc.findall(qn("w:p")):
        txt = "".join(node.text or "" for node in p.iter(qn("w:t"))).strip()
        if txt:
            texts.append(txt)
    return texts


def _make_paragraph(para, in_table, date=""):
    """Build one Paragraph; return None for an empty non-list spacer paragraph."""
    text = para.text.strip()
    is_list = has_numbering(para)
    if not text and not is_list:
        return None
    return Paragraph(
        text=text,
        style_name=para.style.name if para.style else "(none)",
        rendered_size=effective_size(para),
        is_bold=any(r.bold for r in para.runs),
        has_num_pr=is_list,
        in_table=in_table,
        date=date,
        override=any(r.font.size is not None for r in para.runs),
    )


def load_docx(path) -> list[Paragraph]:
    """Flatten a .docx into Paragraph objects: body paragraphs then table cells.

    Empty non-list paragraphs (layout spacers) are dropped.
    """
    doc = Document(str(path))
    paragraphs: list[Paragraph] = []

    for para in doc.paragraphs:
        p = _make_paragraph(para, in_table=False)
        if p is not None:
            paragraphs.append(p)

    for table in doc.tables:
        for tr in table._tbl.tr_lst:
            tcs = tr.tc_lst
            if not tcs:
                continue
            content_cell = _Cell(tcs[0], table)
            dates = _date_cell_texts(tcs[-1]) if len(tcs) > 1 else []
            date_idx = 0
            for para in content_cell.paragraphs:
                text = para.text.strip()
                is_list = has_numbering(para)
                if not text and not is_list:
                    continue
                date = ""
                if not is_list and date_idx < len(dates):
                    date = dates[date_idx]
                    date_idx += 1
                paragraphs.append(_make_paragraph(para, in_table=True, date=date))

    return paragraphs
